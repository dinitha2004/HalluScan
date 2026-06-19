"""Generate docs/HallKing_Overview.docx — a NEWCOMER-FRIENDLY explanation of the project:
what each technique is, how it detects hallucinations (with everyday analogies), what is
novel about combining them, how HallKing works end-to-end, and how we prove it works.

Run in se_probes_env:  python tools/make_overview_doc.py
"""
import os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs", "HallKing_Overview.docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def H(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = NAVY
        return h

    def P(text, bold=False, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        return p

    def bullet(text, bold_lead=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_lead:
            r = p.add_run(bold_lead); r.bold = True
            p.add_run(text)
        else:
            p.add_run(text)
        return p

    # ---------------------------------------------------------------- title
    title = doc.add_heading("HallKing", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("A fused hallucination detector for Large Language Models\n"
                     "combining Semantic Entropy Probes + HalluShift + TSV")
    rs.italic = True; rs.font.size = Pt(12)
    doc.add_paragraph()

    # ---------------------------------------------------------------- 1. problem
    H("1. The problem: when an AI sounds confident but is wrong", 1)
    P("Large Language Models (LLMs) like Llama sometimes produce answers that are fluent, "
      "well-structured, and completely made up. This is called a hallucination. For example, "
      "if you ask “What is the offspring of a male lion and a tigress called?” the model "
      "might confidently answer “a liger” (correct) — or just as confidently invent a wrong "
      "fact. Because the wrong answers look exactly as confident as the right ones, a normal user "
      "cannot tell them apart.")
    P("In low-stakes uses this is annoying; in medicine, law, or finance it is dangerous. A "
      "hallucination detector is a second system that looks at the model and its answer and "
      "estimates: “how likely is this answer to be made up?” HallKing is such a detector.")
    P("The key idea HallKing builds on: the model often “knows” internally when it is unsure, "
      "even while its words sound confident. That hidden uncertainty leaves fingerprints inside "
      "the model’s internal numbers (its “hidden states”) and in the probabilities it assigns to "
      "each word. Different detectors read different fingerprints.")

    # ---------------------------------------------------------------- 2. three lenses
    H("2. Three lenses on the same answer", 1)
    P("HallKing combines three existing, published detectors. Each looks at the answer through a "
      "different “lens.” Think of three doctors examining the same patient with different "
      "instruments — each catches things the others miss.")

    H("2a. Semantic Entropy Probes (SEP) — “Does the model sound unsure of the meaning?”", 2)
    P("Analogy: imagine asking the same person the same question several times. If they truly know "
      "the answer, they say the same thing every time (“Paris… Paris… Paris”). If they are "
      "guessing, the meaning of their answers jumps around (“Rome… Naples… Berlin”). That "
      "variability in meaning is called semantic entropy — high entropy means the model is "
      "uncertain about the meaning, which strongly correlates with hallucination.")
    P("Normally, measuring this requires generating the answer 5–10 times (expensive). SEP’s clever "
      "trick: it trains a tiny, simple classifier (a “probe”) to predict that semantic entropy "
      "directly from the model’s internal hidden states after just ONE answer. So SEP reads the "
      "model’s internal state and outputs “how semantically uncertain is the model right now?” — "
      "cheaply, from a single generation.")
    bullet(" the probe predicts high semantic entropy (or low accuracy) from the hidden states.",
           bold_lead="How it flags a hallucination:")

    H("2b. HalluShift — “Does the model’s train of thought lurch as it answers?”", 2)
    P("Analogy: a confident expert’s reasoning flows smoothly; a bluffer’s reasoning stutters and "
      "jumps around. An LLM processes text through many internal layers, one after another. "
      "HalluShift watches how the internal “picture” changes from layer to layer and how steady the "
      "model’s word-by-word confidence is.")
    P("Concretely it measures: (1) how much the internal representation shifts between nearby layers "
      "(using a distance called the Wasserstein distance) and how similar they stay (cosine "
      "similarity), and (2) statistics of the per-word probabilities (e.g., the least-confident word, "
      "how spread out the confidence is, how sharply it changes). A small neural network turns all "
      "these signals into one hallucination score.")
    bullet(" turbulent, lurching internal dynamics + shaky word probabilities push the score up.",
           bold_lead="How it flags a hallucination:")

    H("2c. TSV (Truthfulness Separator Vector) — “Does the answer point toward the truthful direction?”", 2)
    P("Analogy: imagine a room where truthful answers naturally gather in one corner and made-up "
      "answers in another. In a raw LLM the two groups are jumbled together. TSV learns a single "
      "small “nudge” vector that, when added inside the model, pulls the two groups apart into clean "
      "corners. It also learns the centre point (“prototype”) of each corner.")
    P("To judge a new answer, TSV nudges the model, looks at where the answer lands, and measures "
      "whether it sits closer to the truthful corner or the hallucinated corner. Remarkably, this "
      "nudge is tiny (about 4,000 numbers) and is trained from only a handful of labelled examples.")
    bullet(" the answer’s internal representation sits closer to the hallucinated prototype "
           "than the truthful one.", bold_lead="How it flags a hallucination:")

    # ---------------------------------------------------------------- 3. novelty
    H("3. What is special about HallKing (the novelty)", 1)
    P("Each of the three detectors is published and works on its own. The novelty of HallKing is "
      "combining all three into a single, calibrated detector — and the reason this helps is that "
      "the three lenses are complementary rather than redundant:")
    bullet(" predicted from the model’s internal state (an output-distribution view).",
           bold_lead="SEP measures semantic uncertainty")
    bullet(" across the model’s layers and word probabilities (a process / dynamics view).",
           bold_lead="HalluShift measures internal turbulence")
    bullet(" in a representation space reshaped specifically to separate truth from fiction "
           "(a learned-geometry view).", bold_lead="TSV measures position")
    P("Because they look at different evidence, when one lens is fooled another often is not, so a "
      "combiner that weighs all three should beat any single detector. This is supported by prior "
      "research (the HaMI paper) which found that adding an uncertainty signal on top of internal "
      "representations measurably improves detection. To our knowledge, no prior work fuses these "
      "three specific detectors together.")
    P("Three further practical strengths: (a) all three work from a SINGLE generated answer (no "
      "expensive repeated sampling), so HallKing stays cheap and can run on a single consumer GPU; "
      "(b) the same fused model can highlight WHICH sentence of a long answer is likely "
      "hallucinated, not just give one number; and (c) the whole thing is designed to be hosted as "
      "a live web demo (a Colab backend with a Vercel front-end).")

    # ---------------------------------------------------------------- 4. how it works
    H("4. How HallKing works, end to end", 1)
    P("In words (a single diagram-as-text):")
    for step in [
        ("Question in.", " The user asks a question."),
        ("One answer.", " Llama-3.1-8B generates one concise, most-likely answer."),
        ("Three scores.", " From that single answer, the three detectors each produce a hallucination "
                          "score — SEP and HalluShift read the model’s internals during/after the answer; "
                          "TSV checks the answer’s position in its truth-separated space."),
        ("Fusion.", " A small “meta-classifier” (a simple logistic-regression / gradient-boosting "
                    "model) combines the three scores into ONE calibrated probability of hallucination."),
        ("Output.", " A headline score for the whole answer, plus optional per-sentence highlighting "
                    "so you can see exactly where a long answer goes wrong."),
    ]:
        bullet(step[1], bold_lead=step[0])
    P("A careful engineering detail: TSV physically nudges the model, while SEP and HalluShift need "
      "the un-nudged model. HallKing therefore applies the nudge only for TSV’s reading and removes "
      "it immediately, so the other two always see the original model.")

    # ---------------------------------------------------------------- 5. proof
    H("5. How we prove it works", 1)
    P("We evaluate on hundreds of questions from standard benchmarks (TruthfulQA and TriviaQA). For "
      "each answer we obtain a trustworthy “ground-truth” label using BLEURT, an automatic measure "
      "of how well the answer matches known correct answers (the same labelling all three original "
      "papers use). Then we report, for each individual detector AND for the fused detector:")
    bullet(" the area under the ROC curve — how well scores rank hallucinations above truthful "
           "answers (the headline number the original papers report).", bold_lead="AUROC:")
    bullet(" the area under the precision-recall curve — accuracy when hallucinations are rare.",
           bold_lead="AUPR:")
    bullet(" how many answers were correctly vs. incorrectly flagged at a chosen cut-off, "
           "plus accuracy / precision / recall / F1.", bold_lead="Confusion matrix:")
    P("Success criterion: the fused detector should match or beat the best single detector, and land "
      "in the same range as the numbers reported in the three source papers — giving a direct, "
      "apples-to-apples comparison.")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("Sources: Semantic Entropy Probes (Kossen et al., 2024, arXiv:2406.15927); "
                      "HalluShift (Dasgupta et al., 2025, arXiv:2504.09482); "
                      "TSV / Steer LLM Latents (Park et al., ICML 2025, arXiv:2503.01917).")
    fr.italic = True; fr.font.size = Pt(9)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"[overview] wrote {OUT}")


if __name__ == "__main__":
    build()
